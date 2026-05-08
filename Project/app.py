import os
import uuid
from datetime import datetime
from dotenv import load_dotenv
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify
from flask_cors import CORS
import pyodbc
import re
import pandas as pd
from flask import Flask, render_template, request, redirect, session
import docx
import matplotlib.pyplot as plt
import io
import base64
import seaborn as sns

load_dotenv()

# Environment variables
ACCESS_DB_PATH = os.getenv("ACCESS_DB_PATH", "C:/Liz/Clark/All_Documents/MBASecondSemester/SchoolSubjects/LizProjectRequirement/Project/data/meetings.accdb")
PRODUCT_DB_PATH = os.getenv("PRODUCT_DB_PATH", "C:/Liz/Clark/All_Documents/MBASecondSemester/SchoolSubjects/LizProjectRequirement/Project/Ennoah.accdb")
BASE_URL = os.getenv("BASE_URL", "http://localhost:5000")
SECRET_KEY = os.getenv("SECRET_KEY", "dev")

app = Flask(__name__)
app.secret_key = SECRET_KEY
CORS(app)  # Enable cross-origin requests

# Fake user database example
users = {
    "elizabeth@yahoo.com": {
        "name": "Elizabeth",
        "password": "Mysite@3987"
    }
}

# ---------- Shared DB Connection ----------
def get_conn(db_path):
    return pyodbc.connect(
        fr"Driver={{Microsoft Access Driver (*.mdb, *.accdb)}};DBQ={db_path};"
    )

# ---------- Meeting Scheduler ----------
def create_meeting(title, organizer_email, attendee_email, start_time, end_time, external_link, notes):
    join_token = uuid.uuid4().hex[:16]
    conn = get_conn(ACCESS_DB_PATH)
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO Meetings (Title, OrganizerEmail, AttendeeEmail, StartTime, EndTime, JoinToken, ExternalLink, Notes)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    """, (title, organizer_email, attendee_email, start_time, end_time, join_token, external_link, notes))
    conn.commit()
    cur.close()
    conn.close()
    return join_token

def get_meetings():
    conn = get_conn(ACCESS_DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT ID, Title, OrganizerEmail, AttendeeEmail, StartTime, EndTime, JoinToken, ExternalLink FROM Meetings ORDER BY StartTime DESC")
    rows = cur.fetchall()
    cur.close()
    conn.close()
    return rows

def get_meeting_by_token(token):
    conn = get_conn(ACCESS_DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT ID, Title, OrganizerEmail, AttendeeEmail, StartTime, EndTime, JoinToken, ExternalLink, Notes FROM Meetings WHERE JoinToken = ?", (token,))
    row = cur.fetchone()
    cur.close()
    conn.close()
    return row

def parse_events(text):
    events = []
    current_date = None

    # Split by lines
    lines = text.split("\n")

    for line in lines:
        line = line.strip()

        # Detect date headers like "Monday, March 30th"
        if re.match(r"^[A-Za-z]+,\s+[A-Za-z]+\s+\d{1,2}", line):
            current_date = line
            continue

        # Detect event titles (non-empty lines that are not sponsors or times)
        if line and not line.startswith("Sponsored by") and not re.search(r"\d{1,2}:\d{2}", line):
            event = {
                "date": current_date,
                "title": line,
                "sponsor": "",
                "time": "",
                "location": "",
                "description": ""
            }
            events.append(event)

        # Sponsor
        if line.startswith("Sponsored by"):
            events[-1]["sponsor"] = line

        # Time + location
        if re.search(r"\d{1,2}:\d{2}", line):
            events[-1]["time"] = line

    return events

def create_chart(df, group_col, title):
    revenue = df.groupby(group_col).apply(lambda x: (x["Units Sold"] * x["Unit Price"]).sum())

    plt.figure(figsize=(6,4))
    revenue.plot(kind="bar", color="#ff9900")
    plt.title(title)
    plt.ylabel("Revenue ($)")
    plt.tight_layout()

    img = io.BytesIO()
    plt.savefig(img, format="png")
    img.seek(0)
    return base64.b64encode(img.getvalue()).decode()

@app.route("/events")
def events():
    doc = docx.Document("Events.docx")   # correct filename
    text = "\n".join([p.text for p in doc.paragraphs])

    events = parse_events(text)
    return render_template("events.html", events=events)

@app.route("/", methods=["GET", "POST"])
def signin():
    if request.method == "POST":
        email = request.form["email"]
        password = request.form["password"]

        if email in users and users[email]["password"] == password:
            session["user"] = users[email]
            return redirect("/streaming")   # go to welcome page after login

        return render_template("signin.html", error="Invalid email or password")

    return render_template("signin.html")   # default page is sign-in
@app.route("/logout")
def logout():
    session.clear()
    return redirect("/")
@app.route("/meetinglist")
def index():
    meetings = get_meetings()
    return render_template("meetinglist.html", meetings=meetings, base_url=BASE_URL)
@app.route("/welcome")
def welcome():
    return render_template("welcome.html")
@app.route('/menu')
def menu():
    return render_template('menu.html')
@app.route('/thankyou')
def thankyou():
    return render_template('thankyou.html')
@app.route('/history')
def history():
    return render_template('history.html')
@app.route("/goodfood")
def goodfood():
    return render_template("goodfood.html")

@app.route("/Stories")
def Stories():
    return render_template("Stories.html")
@app.route("/cart")
def cart():
    return render_template("cart.html")
@app.route('/rewards')
def rewards():
    return render_template('rewards.html')

@app.route("/StylishClothing")
def StylishClothing():
    return render_template("StylishClothing.html")

@app.route("/Bags")
def Bags():
    return render_template("Bags.html")

@app.route("/streaming", methods=["GET", "POST"])
def streaming():
    if request.method == "POST":
        # handle form submission
        pass
    return render_template("streaming.html")

@app.route('/careers')
def careers():
    return render_template('careers.html')

@app.route('/account')
def account():
    return render_template('accounts.html')
@app.route("/poems")
def poems():
    return render_template("poems.html")
@app.route("/education")
def education():
    df = pd.read_csv("classes.csv")
    classes = df.to_dict(orient="records")
    return render_template("education.html", classes=classes)
@app.route("/poems_chunk")
def poems_chunk():
    poems = load_poems()
    page = int(request.args.get("page", 1))
    size = 5
    start = (page - 1) * size
    end = start + size
    chunk = poems[start:end]
    has_more = end < len(poems)
    return jsonify({"poems": chunk, "has_more": has_more})

def load_poems():
    doc = docx.Document("poems.docx")
    poems = []
    current_title = None
    current_text = []

    for para in doc.paragraphs:
        text = para.text.strip()

        if not text:
            if current_title and current_text:
                poems.append({
                    "title": current_title,
                    "text": " ".join(current_text)
                })
            current_title = None
            current_text = []
            continue

        if current_title is None:
            current_title = text
        else:
            current_text.append(text)

    if current_title and current_text:
        poems.append({
            "title": current_title,
            "text": " ".join(current_text)
        })

    return poems

@app.route("/calendar")
def calendar():
    return render_template("calendar.html")

@app.route("/calendar_events")
def calendar_events():
    events = []
    return jsonify(events)

@app.route("/reports")
def reports():
    df = pd.read_csv("sales_data.csv")
    chart_location = create_chart(df, "Location", "Revenue by Location")
    chart_product = create_chart(df, "Product", "Revenue by Product")


   # Basic analytics
    total_revenue = (df["Units Sold"] * df["Unit Price"]).sum()
    revenue_by_location = df.groupby("Location").apply(lambda x: (x["Units Sold"] * x["Unit Price"]).sum())
    revenue_by_product = df.groupby("Product").apply(lambda x: (x["Units Sold"] * x["Unit Price"]).sum())

    df["Revenue"] = df["Units Sold"] * df["Unit Price"]
    daily = df.groupby("Date")["Revenue"].sum()

    plt.figure(figsize=(6,4))
    daily.plot(kind="line", marker="o", color="#ff9900")
    plt.title("Revenue Over Time")
    plt.ylabel("Revenue ($)")
    plt.tight_layout()

    img = io.BytesIO()
    plt.savefig(img, format="png")
    img.seek(0)
    chart_time = base64.b64encode(img.getvalue()).decode()

    df["Cost"] = df["Unit Price"] * 0.6   # assume 40% margin
    df["Profit"] = df["Units Sold"] * (df["Unit Price"] - df["Cost"])
    total_profit = df["Profit"].sum()
    total_profit = round(total_profit, 2) 
    profit_by_product = df.groupby("Product")["Profit"].sum()
    best_sellers = df.groupby("Product")["Units Sold"].sum().sort_values(ascending=False)

    pivot = df.pivot_table(values="Revenue", index="Location", columns="Product", aggfunc="sum")

    plt.figure(figsize=(6,4))
    sns.heatmap(pivot, annot=True, cmap="Oranges")
    plt.title("Revenue Heatmap")
    plt.tight_layout()

    img = io.BytesIO()
    plt.savefig(img, format="png")
    img.seek(0)
    heatmap = base64.b64encode(img.getvalue()).decode()
    daily_ma = daily.rolling(window=3).mean()
    peak_day = daily.idxmax()
    slow_day = daily.idxmin()

    return render_template(
        "reports.html",
        chart_location=chart_location,
        chart_product=chart_product,
        chart_time = chart_time,
        total_revenue=total_revenue,
	total_profit = total_profit,
	best_sellers = best_sellers,
	heatmap = heatmap,
        peak_day = peak_day ,
        slow_day = slow_day,
        revenue_by_location=revenue_by_location,
        revenue_by_product=revenue_by_product
    )
@app.route('/payments')
def payments():
    return render_template('payments.html')

@app.route("/schedule", methods=["GET", "POST"])
def schedule():
    if request.method == "POST":
        title = request.form.get("title") or "Meeting"
        organizer_email = request.form.get("organizer_email") or ""
        attendee_email = request.form.get("attendee_email") or ""
        start_time_str = request.form.get("start_time")
        end_time_str = request.form.get("end_time")
        external_link = request.form.get("external_link") or ""
        notes = request.form.get("notes") or ""

        try:
            start_time = datetime.fromisoformat(start_time_str)
            end_time = datetime.fromisoformat(end_time_str)
            if end_time <= start_time:
                flash("End time must be after start time.", "error")
                return redirect(url_for("schedule"))
        except Exception:
            flash("Invalid date/time format. Use the picker.", "error")
            return redirect(url_for("schedule"))

        token = create_meeting(title, organizer_email, attendee_email, start_time, end_time, external_link, notes)
        join_url = f"{BASE_URL}/join/{token}"
        flash(f"Meeting scheduled. Share this link: {join_url}", "success")
        return redirect(url_for("index"))
    return render_template("schedule.html")

@app.route("/join/<token>")
def join(token):
    mtg = get_meeting_by_token(token)
    if not mtg:
        return render_template("join.html", error="Invalid or expired link.", meeting=None)

    meeting = {
        "id": mtg[0],
        "title": mtg[1],
        "organizer_email": mtg[2],
        "attendee_email": mtg[3],
        "start_time": mtg[4],
        "end_time": mtg[5],
        "token": mtg[6],
        "external_link": mtg[7] or "",
        "notes": mtg[8] or ""
    }
    return render_template("join.html", meeting=meeting)

# ---------- Product API ----------
@app.route("/get-products")
def get_products():
    conn = get_conn(PRODUCT_DB_PATH)
    query = "SELECT [Product Name], Price, Category FROM Products ORDER BY Price;"
    df = pd.read_sql(query, conn)
    conn.close()
    products = df.to_dict(orient="records")
    return jsonify(products)

# ---------- ICS Generator ----------
def build_ics(meeting, base_url):
    return f"""BEGIN:VCALENDAR
VERSION:2.0
PRODID:-//MeetingScheduler//EN
BEGIN:VEVENT
UID:{meeting['token']}@meetings
DTSTAMP:{meeting['start_time'].strftime('%Y%m%dT%H%M%SZ')}
DTSTART:{meeting['start_time'].strftime('%Y%m%dT%H%M%SZ')}
DTEND:{meeting['end_time'].strftime('%Y%m%dT%H%M%SZ')}
SUMMARY:{meeting['title']}
DESCRIPTION:Join link: {base_url}/join/{meeting['token']}
END:VEVENT
END:VCALENDAR
"""

if __name__ == "__main__":
    app.run(debug=True)
